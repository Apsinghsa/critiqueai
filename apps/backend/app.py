from flask import Flask, request, jsonify
from flask_cors import CORS
import pymupdf
import pdfplumber
import os
import docx
import markdown2
from google import genai
from google.genai import types
from datetime import datetime, timezone, timedelta
from PIL import Image
from io import BytesIO
import zipfile
import uuid
from dotenv import load_dotenv
import io
import base64
from supabase import create_client, Client


# Load environment variables from .env file
load_dotenv()


# Updated prompts for AI evaluation
EVALUATION_PROMPT = """Evaluate the following question and answer pair for accuracy and relevance. Write a concise summary of the answer (60–70 words max) in clear, human-like legal language. Then, justify the score with specific reasoning and suggest precise improvements. Do not include any introduction, heading, or reference to the question or answer. Avoid words like 'evaluation', 'response', or similar. Just give the evaluation exactly as instructed. Don't give any score here.
"""

SCORE_PROMPT = """"Evaluate the following question and answer pair. Based on the accuracy and relevance of the answer, assign a single-digit score from 0 to 10 (e.g., '3', '7', '10'). Respond with only the numeric score. 

If the answer is completely incorrect, or if the input is malformed, unclear, or missing either the question or the answer, always return '0'."
"""

SUMMARY_PROMPT = '''Generate a structured summary of the provided data. The data consists of multiple files, each separated by a line of asterisks (*************************).
DONT ADD ANY SUMMARY STARTING LINE. JUST START SUMMART DIRECTLY.

For each file:
Extract key points and organize them under a relevant heading based on the file's content.
Maintain the summary length in proportion to the content's significance; do not unnecessarily shorten it.
Use bullet points for clarity and readability.
Do not include any greetings or salutations.

IMP: Format the output for easy readability.
'''

# New prompt for roadmap generation
ROADMAP_PROMPT = '''Generate a concise roadmap for learning {topic} over 4-6 weeks. The roadmap should be divided into logical sections and include:

Key Concepts – Outline 3-5 foundational concepts or principles in {topic}, each explained in 15-20 words.
Practice Tasks – Provide 2-3 hands-on exercises or projects that reinforce the concepts. Each task should have clear instructions and expected outcomes.
Best Resources – List 3-5 effective resources (books, courses, websites, tools) with clickable links and a one-line description for each.
Format the output for easy readability.'''

IMG_PROMPT = "Describe the content of the image extracted from the PDF."

CONTENT_PROMPT = """
You are a professional educator. Generate well-structured study notes for the topic: '{topic}'.
The student is in '{academic_level}'and prefers '{detail_level}' notes.
Format the notes in '{format_preference}' style.
If the topic is technical, include code snippets: '{technical_content}'.
Adjust content depth, complexity, and length based on these inputs.

Don't add any greeting and thank you note.
"""


# Initialize Flask app
app = Flask(__name__)
# app.secret_key = os.urandom(24)
CORS(app, origins=["http://localhost:3000", "http://localhost:8000"])

# Configure the Google Generative AI API
client = genai.Client(api_key=os.getenv('GEN_API2'))
FLASH = 'gemini-2.0-flash'
FLASH_LITE = 'gemini-2.0-flash-lite'

# Initialize Supabase client
SUPABASE_URL = os.getenv('SUPABASE_URL')
SUPABASE_KEY = os.getenv('SUPABASE_SERVICE_KEY')
if not SUPABASE_URL or not SUPABASE_KEY:
    raise RuntimeError("SUPABASE_URL and SUPABASE_SERVICE_KEY must be set")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)


# Define allowed file extensions
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'txt', 'pdf', 'docx'}


# Function to check allowed file extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def resolve_session(data):
    uid = data.get("uid")
    guest_id = data.get("guestId")

    print("Inside resolve session function and printing the IDs:\n")
    print(uid)
    print(guest_id)

    if uid:
        identifier = uid
        session_id = f"{identifier}_{uuid.uuid4().hex[:8]}"
    elif guest_id:
        identifier = f"guest_{guest_id}"
        session_id = f"{identifier}_{uuid.uuid4().hex[:8]}"
    else:
        raise ValueError("Neither UID nor guestId provided")

    print("Session ID: ", session_id)

    return identifier, session_id


# Function for summary page
def get_evaluation(text, isInstruction=None):
    if isInstruction:  # If an instruction is provided, append it
        full_prompt = f"{SUMMARY_PROMPT}\n\nAdditional Instructions: {isInstruction}"
    else:
        full_prompt = SUMMARY_PROMPT

    response = client.models.generate_content(
        model=FLASH,
        contents=[text, full_prompt],
        config=types.GenerateContentConfig(
            max_output_tokens=1000,  # Set the token limit back to 1500
            temperature=0.5)
    )

    content = response.text if response.candidates else "No summary generated"
    return content


# <----------------- Evaluation Page Here Starts From Here ------------------->
# Functions for evaluation page
def get_evaluate(text):
    response = client.models.generate_content(
        model=FLASH,
        contents=[text, EVALUATION_PROMPT],
        config=types.GenerateContentConfig(
            max_output_tokens=1000,
            temperature=0.5))
    score = client.models.generate_content(
        model=FLASH,
        contents=[text, SCORE_PROMPT],
        config=types.GenerateContentConfig(
            max_output_tokens=1000,
            temperature=0.5))
    return response, score


# Text File
def process_txt_file(content):
    return get_evaluate(content)


# PDF File
def process_pdf_file(pdf_bytes):
    response = client.models.generate_content(
        model=FLASH,
        contents=[types.Part.from_bytes(
            data=pdf_bytes,
            mime_type='application/pdf'),
            EVALUATION_PROMPT],
        config=types.GenerateContentConfig(
            max_output_tokens=200,
            temperature=0.5,
        )
    )
    score = client.models.generate_content(
        model=FLASH,
        contents=[types.Part.from_bytes(
            data=pdf_bytes,
            mime_type='application/pdf'),
            SCORE_PROMPT],
        config=types.GenerateContentConfig(
            max_output_tokens=200,
            temperature=0.5,
        )
    )
    return response, score

# Image File
def process_img_file(image_content):
    img_base64 = base64.b64encode(image_content).decode("utf-8")
    response = client.models.generate_content(
        model=FLASH,
        contents=[
            {"text": "Extract all text from this image. Return only the extracted text, nothing else."},
            {"inline_data": {
                "mime_type": "image/png",
                "data": img_base64
            }}
        ],
        config=types.GenerateContentConfig(
            max_output_tokens=2000,
            temperature=0.1
        )
    )
    extracted_text = response.text if response.candidates else ""
    return get_evaluate(extracted_text)


# Docx File
def process_docx_file(doc_bytes):
    doc = docx.Document(BytesIO(doc_bytes))
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    docx_text = '\n'.join(full_text)

    return get_evaluate(docx_text)

# <----------------- Evaluate Page Functions Ends Here ------------------>


# <----------------- Helper functions for summary_out page ------------------>
# Image Extraction for PDF
def describe_image_with_gemini(image, flag=False):
    """
    Process image using Gemini with proper encoding and fallback
    Returns: Gemini description (preferred) or Vision API analysis (fallback)
    """

    # Encode to base64
    img_base64 = encode_image(image)

    # Generate description with Gemini
    response = client.models.generate_content(
        model=FLASH_LITE,
        contents=[
            {"text": IMG_PROMPT},
            {"inline_data": {
                "mime_type": "image/png",
                "data": img_base64
            }}
        ],
        config=types.GenerateContentConfig(
            temperature=0.3,
            max_output_tokens=1000
        )
    )

    # Return formatted response
    if response.candidates:
        return response.text.strip()
    return "No description generated"


# Encoding image for PDF
def encode_image(image):
    """Convert PIL Image to Base64 format for Gemini."""
    if isinstance(image, bytes):
        image = Image.open(io.BytesIO(image))
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")


# Image Extraction for PDF
def extract_images(pdf_file, page_index):
    """Extracts images from a PDF page with Gemini descriptions."""
    page = pdf_file.load_page(page_index)
    image_list = page.get_images(full=True)
    images = []

    for img_index, img in enumerate(image_list, start=1):
        try:
            xref = img[0]
            base_image = pdf_file.extract_image(xref)
            image_bytes = base_image["image"]  # Raw bytes
            image_ext = base_image["ext"]
            image_name = f"image_{page_index+1}_{img_index}.{image_ext}"

            image = Image.open(io.BytesIO(image_bytes))

            # Generate description using Gemini with raw bytes
            image_description = describe_image_with_gemini(image)

            images.append({
                "name": image_name,
                "description": image_description
            })
        except Exception as e:
            images.append(f"[Image extraction error: {str(e)}]")
    return images


# Text Extraction for PDF
def extract_text(pdf_file, page_index):
    """Extracts text from a given PDF page."""
    page = pdf_file.load_page(page_index)
    return page.get_text("text")


# Tables Extraction for PDF
def extract_tables(content, page_index):
    """Extracts tables using pdfplumber."""
    tables = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            if page_index < len(pdf.pages):
                page = pdf.pages[page_index]
                extracted_tables = page.extract_tables()
                for table_index, table in enumerate(extracted_tables, start=1):
                    tables.append({
                        "table_number": table_index,
                        "data": table
                    })
    except Exception as e:
        tables.append(f"[Table error: {str(e)}]")
    return tables


# Process PDF for Page-Wise Extraction
def extract_pdf_content(pdf_bytes):
    """Extracts text, images, and tables from PDF bytes and returns as a string."""
    pdf_file = pymupdf.open(
        stream=pdf_bytes, filetype="pdf")  # Open from bytes
    extracted_data = []

    for page_index in range(len(pdf_file)):
        page_content = {
            "page_number": page_index + 1,
            "text": extract_text(pdf_file, page_index),
            "images": extract_images(pdf_file, page_index),
            # Directly pass bytes
            "tables": extract_tables(pdf_bytes, page_index)
        }
        extracted_data.append(page_content)

    # Convert the extracted_data list to a string
    result = "PDF CONTENT BELOW HERE\n"
    for page in extracted_data:
        result += f"PAGE {page['page_number']}:\n"
        result += f"TEXT:\n{page['text']}\n"

        result += "IMAGES:\n"
        for img in page['images']:
            if isinstance(img, dict):
                result += f"--> {img['name']}: {img['description']}\n"
            else:
                result += f"--> {img}\n"  # Error message case
        result += "\n"

        result += "TABLES:\n"
        for table in page['tables']:
            if isinstance(table, dict):
                result += f"Table {table['table_number']}:\n"
                for row in table['data']:
                    result += f"{row}\n"
                result += "\n"
            else:
                result += f"--> {table}\n"  # Error message case
        result += "---------------------------------------------------\n"

    return result
# PDF FUNCTION ENDS HERE


# Functions to read .docx files
def read_docx_file(docx_content):
    """Read the content of a DOCX file and return it as a string, including tables."""
    doc = docx.Document(BytesIO(docx_content))
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            full_text.append('\t'.join(row_data))

    return '\n'.join(full_text)


def extract_docx_content(content):
    try:
        # Extract regular text content
        text_content = read_docx_file(content)

        # Extract and process images
        image_texts = []
        with zipfile.ZipFile(BytesIO(content)) as zip_file:
            for file_info in zip_file.infolist():
                if file_info.filename.startswith('word/media/'):
                    image_data = zip_file.read(file_info)
                    try:
                        img_text = describe_image_with_gemini(image_data)
                        image_texts.append(img_text)
                    except Exception as e:
                        # Handle image processing errors silently or log them
                        pass

        # Combine all text elements
        combined_text = f"{text_content}\n{' '.join(image_texts)}"

        return combined_text

    except Exception as e:
        return f"DOCX Error: {str(e)}"


def upload_files(file, session):
    try:
        unique_filename = f"{uuid.uuid4()}.{file.filename.split('.')[-1]}"
        file_bytes = file.read()
        supabase.storage.from_('user-files').upload(
            f"sessions/{session}/{unique_filename}",
            file_bytes,
            {"content-type": file.content_type}
        )
        url = supabase.storage.from_('user-files').get_public_url(
            f"sessions/{session}/{unique_filename}"
        )
        supabase.table('user_files').insert({
            "session_id": session,
            "filename": unique_filename,
            "url": url,
            "uploaded_at": datetime.now(timezone.utc).isoformat()
        }).execute()
        return True
    except Exception as e:
        print(f"Upload failed for {file.filename}: {e}")
        return False

# Function to get file content from Google Cloud Storage
def get_user_files(session_id, time):
    """Fetches URLs for a user from Supabase."""

    response = supabase.table('user_files').select('*').eq(
        'session_id', session_id
    ).execute()
    file_urls = []
    for record in response.data:
        if record['uploaded_at'] >= time.isoformat():
            file_urls.append(record['url'])

    return file_urls


def read_file_from_storage(file_url):
    """Reads a file as bytes directly from Supabase Storage."""
    file_path = file_url.split('/storage/v1/object/public/user-files/')[-1]
    file_bytes = supabase.storage.from_('user-files').download(file_path)
    filename = file_path.split('/')[-1]
    return file_bytes, filename


def process_file(file_url):
    """Processes a file based on its type (Text, PDF, DOCX, Image)."""
    file_bytes, filename = read_file_from_storage(file_url)
    file_ext = filename.split(".")[-1].lower()

    # Plain text files
    if file_ext in ["txt", "csv", "json"]:
        return file_bytes.decode("utf-8")

    # PDF files (Text + Images)
    elif file_ext == "pdf":
        final_extracted_content = extract_pdf_content(file_bytes)
        return final_extracted_content

    # Image processing
    elif file_ext in ["png", "jpg", "jpeg"]:
        return describe_image_with_gemini(file_bytes, True)

    # Docx Processing
    elif file_ext == "docx":
        return extract_docx_content(file_bytes)
    else:
        raise ValueError("Unsupported file type")
# <---------- Summary Page Functions Ends Here ------------>


# <---------- Save Output to Supabase ----------->
def set_output_helper(user, head, data, time, typ):
    output_id = str(uuid.uuid4())
    supabase.table('outputs').insert({
        "id": output_id,
        "user_id": user,
        "name": head,
        "time": time,
        "content": data,
        "type": typ
    }).execute()
    return output_id


def get_output_helper(user):
    response = supabase.table('outputs').select('*').eq(
        'user_id', user
    ).execute()

    output_list = []

    for doc in response.data:
        output_list.append({
            "id": doc['id'],
            "name": doc.get('name'),
            "type": doc.get('type'),
            "time": doc.get('time'),
            "content": doc.get('content')
        })

    return output_list


@app.route("/")
def home():
    # print("home")
    return 'Hello Programmer, Visit our <a href="https://github.com/ashokkmt/critiqueai" target="_blank">GitHub</a>!'


@app.route('/content-out', methods=['POST'])
def generate_content():
    try:
        if request.method == 'POST':
            data = request.get_json()
            if not data:
                return "Please enter a valid topic", 400

            topic = data.get('topic')
            academic_level = data.get('academicLevel')
            detail_level = data.get('detailLevel')
            format = data.get('format')
            include_code = data.get('includeCode')

            # Format the prompt
            full_prompt = CONTENT_PROMPT.format(
                topic=topic,
                academic_level=academic_level,
                detail_level=detail_level,
                format_preference=format,
                technical_content=include_code,
            )

            try:
                print(full_prompt)
                response = client.models.generate_content(
                    model=FLASH,
                    contents=full_prompt,
                    config=types.GenerateContentConfig(
                        max_output_tokens=8192,
                        temperature=0.5
                    )
                )

                temp = markdown2.markdown(
                    response.text,
                    extras=["fenced-code-blocks", "tables",
                            "strike", "task_list", "header-ids"]
                )
                return jsonify({"output": temp}), 200

            except Exception as e:
                return jsonify({"error": f"Error Generating Notes: {str(e)}"}), 500
    except Exception as e:
        print("ERROR: ", str(e))
        return jsonify({"error": str(e)}), 500


# New route for generating roadmap
@app.route('/get-roadmap', methods=['POST'])
def get_roadmap():
    if request.method == 'POST':
        data = request.data.decode('utf-8')
        if not data:
            return "Please enter a valid topic", 400

        print("Topic Received: ", data)
        try:
            full_prompt = ROADMAP_PROMPT.format(topic=data)
            response = client.models.generate_content(
                model=FLASH,
                contents=full_prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=2500,
                    temperature=0.5
                )
            )
            temp = markdown2.markdown(
                response.text,
                extras=["fenced-code-blocks", "tables",
                        "strike", "task_list", "header-ids"]
            )
            if response.candidates:
                return jsonify({"output": temp}), 200
            else:
                return jsonify({"error": "Failed to generate roadmap"}), 500

        except Exception as e:
            return jsonify({"error": f"Error generating roadmap: {str(e)}"}), 500


# Route for Summary output page
@app.route('/summary-out', methods=['POST'])
def summary_out():
    if request.method == 'POST':
        check_file = 'files' in request.files and request.files.getlist(
            'files')
        check_fname = 'text' in request.form and request.form['text']
        combined_text = ""
        time = datetime.now(timezone.utc)

        data = request.form
        guest_id = request.form.get('guestId')
        print("Inside summary Function: ", guest_id, "\n")

        user_id, session_id = resolve_session(data)
        try:
            if check_file and check_fname:
                instruction = request.form.get('text')
                files = request.files.getlist('files')

                for f in files:
                    if f and allowed_file(f.filename):
                        upload_files(f, session_id)
                    else:
                        combined_text += f"Invalid or unsupported file: {f.filename}\n"

                # session_id = session.get('session_id')
                file_url = get_user_files(session_id, time)
                for url in file_url:
                    combined_text += process_file(
                        url) + "\n\n***************************************************\n\n"

                print("Instructions - ", instruction)
                combined_summary = get_evaluation(
                    combined_text, isInstruction=instruction)
                temp = markdown2.markdown(
                    combined_summary,
                    extras=["fenced-code-blocks", "tables",
                            "strike", "task_list", "header-ids"]
                )
                return jsonify({"output": temp}), 200
            elif check_file:
                files = request.files.getlist('files')
                for f in files:
                    if f and allowed_file(f.filename):
                        upload_files(f, session_id)
                    else:
                        combined_text += f"Invalid or unsupported file: {f.filename}\n"

                # session_id = session.get('session_id')
                file_url = get_user_files(session_id, time)
                for url in file_url:
                    combined_text += process_file(
                        url) + "\n\n***************************************************\n\n"

                combined_summary = get_evaluation(combined_text)
                temp = markdown2.markdown(
                    combined_summary,
                    extras=["fenced-code-blocks", "tables",
                            "strike", "task_list", "header-ids"]
                )
                return jsonify({"output": temp}), 200
            else:
                return jsonify({"error": "Invalid input received. No files or text provided."}), 400
        except Exception as e:
            return jsonify({"error": f"Server exception: {str(e)}"}), 500


@app.route('/set-output', methods=['POST'])
def set_output():
    print("nested collection function started")
    data = request.get_json()
    user = data.get('uid')
    heading = data.get('heading')
    time = data.get('time')
    content = data.get('content')
    tp = data.get('type')

    res = set_output_helper(user, heading, content, time, tp)
    print("Output Id - ", res)
    return res


@app.route('/get-output', methods=['POST'])
def get_output():
    data = request.get_json()
    user = data.get("uid")
    if not user:
        return jsonify({"error": "User ID is required"}), 400

    try:
        outputs = get_output_helper(user)
        return jsonify({"outputs": outputs})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/view', methods=['POST'])
def view_output():
    try:
        data = request.json
        uid = data.get('uid')
        doc_id = data.get('doc_id')

        if not uid or not doc_id:
            return jsonify({"error": "Missing UID or Document ID"}), 400

        response = supabase.table('outputs').select('*').eq(
            'id', doc_id
        ).eq('user_id', uid).execute()

        if response.data:
            return jsonify(response.data[0]), 200
        else:
            return jsonify({"error": "Document not found"}), 404

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/delete-output", methods=['DELETE'])
def delete_output():
    data = request.get_json()
    user_id = data.get("user_id")
    doc_id = data.get("doc_id")

    if not user_id or not doc_id:
        return jsonify({"error": "Missing user_id or doc_id"}), 400

    try:
        supabase.table('outputs').delete().eq(
            'id', doc_id
        ).eq('user_id', user_id).execute()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/share-output', methods=['POST'])
def share_output():
    data = request.get_json()
    doc_id = data.get("doc_id")
    user_id = data.get("user_id")

    if not doc_id or not user_id:
        return jsonify({"error": "outputDocId and userUid required"}), 400

    original = supabase.table('outputs').select('*').eq(
        'id', doc_id
    ).eq('user_id', user_id).execute()

    if not original.data:
        return jsonify({"error": "Original output document not found"}), 404

    original_data = original.data[0]
    original_data['copyTime'] = datetime.now(timezone.utc).isoformat()
    shared_id = str(uuid.uuid4())
    original_data['id'] = shared_id
    supabase.table('shared').insert(original_data).execute()

    return jsonify({"sharedDocId": shared_id}), 200


@app.route('/shared/<doc_id>', methods=['GET'])
def get_shared_doc(doc_id):
    try:
        response = supabase.table('shared').select('*').eq('id', doc_id).execute()

        if response.data:
            return jsonify(response.data[0]), 200
        else:
            return jsonify({"error": "Shared document not found"}), 404

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/delete-shared-docs', methods=['GET'])
def delete_expired_shared_docs():
    print("Cleanup for shared docs started.")

    expiration_time = (datetime.now(timezone.utc) - timedelta(minutes=30)).isoformat()

    supabase.table('shared').delete().lt('copyTime', expiration_time).execute()

    return jsonify({"status": "cleaned"}), 200


# Route for deleting user files from Supabase bucket
@app.route('/delete-user-files', methods=['GET'])
def accountcleanup():
    print("Delete Function is called.")
    expiration_time = (datetime.now(timezone.utc) - timedelta(minutes=2)).isoformat()

    docs = supabase.table('user_files').select('*').lt('uploaded_at', expiration_time).execute()

    deleted_files = []

    for doc in docs.data:
        session_id = doc.get('session_id')
        filename = doc.get('filename')

        if session_id and filename:
            file_path = f"sessions/{session_id}/{filename}"
            try:
                supabase.storage.from_('user-files').remove(file_path)
                supabase.table('user_files').delete().eq('id', doc['id']).execute()
                print(f"Deleted: {file_path}")
                deleted_files.append(filename)
            except Exception as e:
                print(f"Failed to delete {file_path}: {e}")

    if not deleted_files:
        return "No Files\n", 200

    return jsonify({"deleted_files": deleted_files}), 200

# Route for evalation page
@app.route('/evaluate', methods=['POST'])
def evaluate():
    try:
        if request.method == 'POST':
            if 'file' in request.files:
                data = request.form
                guest_id = request.form.get('guestId')

                print("Inside Evaluation Function: ", guest_id, "\n")

                user_id, session_id = resolve_session(data)
                file = request.files['file']
                if file.filename == '':
                    return jsonify({"error": "No file selected"}), 400

                if file and allowed_file(file.filename):
                    time = datetime.now(timezone.utc)
                    upload_files(file, session_id)
                    # session_id = session.get('session_id')
                    file_url = get_user_files(session_id, time)
                    file_content, name = read_file_from_storage(file_url[0])
                    file_extension = name.split(".")[-1].lower()

                    if file_extension == 'txt':
                        response, score_response = process_txt_file(
                            file_content.decode('utf-8'))
                    elif file_extension == 'pdf':
                        response, score_response = process_pdf_file(
                            file_content)
                    elif file_extension in {'png', 'jpg', 'jpeg'}:
                        response, score_response = process_img_file(
                            file_content)
                    elif file_extension == 'docx':
                        response, score_response = process_docx_file(
                            file_content)
                    else:
                        return jsonify({"error": "Unsupported file type"}), 400

                    print("response-", response.text,
                          "score-", score_response.text)
                    evaluation_html = markdown2.markdown(
                        response.text,
                        extras=["fenced-code-blocks", "tables",
                                "strike", "task_list", "header-ids"]
                    )
                    return jsonify({"evaluation": evaluation_html, "score": score_response.text}), 200

                else:
                    return jsonify({"error": "Invalid file"}), 400
            elif request.is_json:
                data = request.get_json()
                text = data.get('text')
                if not text:
                    return jsonify({"error": "No text provided"}), 400
                response, score_response = get_evaluate(text)
                print("response-", response.text,
                      " score-", score_response.text)
                evaluation_md = response.text
                evaluation_html = markdown2.markdown(
                    evaluation_md,
                    extras=["fenced-code-blocks", "tables",
                            "strike", "task_list", "header-ids"]
                )
                return jsonify({"evaluation": evaluation_html, "score": score_response.text}), 200
            else:
                return jsonify({"error": "No valid input provided"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# Run the app in debug mode for development
if __name__ == '__main__':
    app.run(debug=True)
